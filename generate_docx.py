from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def parse_memo_to_sections(memo_text: str) -> dict:
    sections = {
        "executive_summary": "",
        "market_analysis": "",
        "team_assessment": "",
        "financial_analysis": "",
        "risk_assessment": "",
        "recommendation": "",
        "company_name": "Company",
        "stage": "",
        "sector": "",
    }
    lines = memo_text.strip().split('\n')
    current_section = None
    buffer = []

    def flush(sec):
        if sec and buffer:
            sections[sec] = '\n'.join(buffer).strip()

    section_map = {
        'executive summary': 'executive_summary',
        'market': 'market_analysis',
        'team': 'team_assessment',
        'financial': 'financial_analysis',
        'risk': 'risk_assessment',
        'recommendation': 'recommendation',
        'investment recommendation': 'recommendation',
    }

    for line in lines:
        stripped = line.strip().lstrip('#').strip()
        lower = stripped.lower()

        matched = False
        for key, sec in section_map.items():
            if key in lower and len(stripped) < 60:
                flush(current_section)
                buffer = []
                current_section = sec
                matched = True
                break

        if not matched:
            if current_section:
                buffer.append(line)
            elif stripped:
                if 'company' in lower or stripped.startswith('#'):
                    clean = stripped.lstrip('#').strip()
                    if len(clean) < 60 and clean:
                        sections['company_name'] = clean

    flush(current_section)
    return sections

def generate_memo_docx(memo_text: str, company_name: str = None) -> bytes:
    sections = parse_memo_to_sections(memo_text)
    if company_name:
        sections['company_name'] = company_name

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header — company name
    header_p = doc.add_paragraph()
    header_run = header_p.add_run(sections['company_name'])
    header_run.bold = True
    header_run.font.size = Pt(18)
    header_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    header_p.paragraph_format.space_after = Pt(2)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Investment Memo — Confidential")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub_run.italic = True
    sub_p.paragraph_format.space_after = Pt(12)

    add_horizontal_rule(doc)

    # Executive Summary
    add_section_heading(doc, "Executive Summary")
    text = sections.get('executive_summary', '')
    for para in text.split('\n'):
        para = para.strip()
        if not para:
            continue
        if para.startswith('- ') or para.startswith('• '):
            add_bullet(doc, para.lstrip('-•').strip())
        else:
            add_body(doc, para)

    add_horizontal_rule(doc)

    # Market Analysis
    add_section_heading(doc, "Market Analysis")
    text = sections.get('market_analysis', '')
    for para in text.split('\n'):
        para = para.strip()
        if not para:
            continue
        if para.startswith('- ') or para.startswith('• '):
            add_bullet(doc, para.lstrip('-•').strip())
        else:
            add_body(doc, para)

    add_horizontal_rule(doc)

    # Team Assessment
    add_section_heading(doc, "Team Assessment")
    text = sections.get('team_assessment', '')
    for para in text.split('\n'):
        para = para.strip()
        if not para:
            continue
        if para.startswith('- ') or para.startswith('• '):
            add_bullet(doc, para.lstrip('-•').strip())
        else:
            add_body(doc, para)

    add_horizontal_rule(doc)

    # Financial Analysis
    add_section_heading(doc, "Financial Analysis")
    text = sections.get('financial_analysis', '')
    for para in text.split('\n'):
        para = para.strip()
        if not para:
            continue
        if para.startswith('- ') or para.startswith('• '):
            add_bullet(doc, para.lstrip('-•').strip())
        else:
            add_body(doc, para)

    add_horizontal_rule(doc)

    # Risk Assessment
    add_section_heading(doc, "Risk Assessment")
    text = sections.get('risk_assessment', '')
    for para in text.split('\n'):
        para = para.strip()
        if not para:
            continue
        if para.startswith('- ') or para.startswith('• '):
            add_bullet(doc, para.lstrip('-•').strip())
        else:
            add_body(doc, para)

    add_horizontal_rule(doc)

    # Investment Recommendation
    add_section_heading(doc, "Investment Recommendation")
    text = sections.get('recommendation', '')
    for para in text.split('\n'):
        para = para.strip()
        if not para:
            continue
        if para.startswith('- ') or para.startswith('• '):
            add_bullet(doc, para.lstrip('-•').strip())
        else:
            add_body(doc, para)

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run("Generated by VC Memo Agent · Confidential · Not for distribution")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    footer_run.italic = True
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
